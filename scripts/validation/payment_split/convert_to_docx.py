from docx import Document
import os

# Create a new Document
document = Document()

# --------------------------------------------------
# Title and Introduction
# --------------------------------------------------
document.add_heading('Income Transfer Workflow for OpenDental', level=1)
document.add_heading('Standard Operating Procedure', level=2)

intro = (
    "This document outlines the proper workflow for identifying, executing, "
    "and verifying income transfers in OpenDental. Proper adherence to this procedure "
    "will reduce unassigned provider transactions and improve financial reporting accuracy."
)
document.add_paragraph(intro)

# --------------------------------------------------
# Section 1: Identifying When Income Transfers Are Needed
# --------------------------------------------------
document.add_heading('1. Identifying When Income Transfers Are Needed', level=2)

# Scenarios Requiring Income Transfer Table
document.add_heading('Scenarios Requiring Income Transfer', level=3)
scenarios = [
    ["Scenario", "Description", "Priority"],
    ["Initial Misallocation", "Payment was assigned to incorrect provider at time of entry", "High"],
    ["Unassigned Payment", "Payment was entered without a provider assignment", "High"],
    ["Prepayment Allocation", "Prepayment needs to be allocated to provider who performed work", "Medium"],
    ["Provider Transfer", "Patient transferred between providers during treatment", "Medium"],
    ["Treatment Plan Change", "Original provider differs from provider who performed service", "Medium"],
    ["Provider Left Practice", "Redistributing payments from departing provider", "Low"],
]
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = scenarios[0][0]
hdr_cells[1].text = scenarios[0][1]
hdr_cells[2].text = scenarios[0][2]
for row in scenarios[1:]:
    row_cells = table.add_row().cells
    row_cells[0].text = row[0]
    row_cells[1].text = row[1]
    row_cells[2].text = row[2]

# Key Indicators (as bullet points)
document.add_heading('Key Indicators', level=3)
indicators = [
    "Provider production report shows discrepancies",
    "Unassigned provider appears in transaction reports",
    "Payment split report shows provider ≠ procedure provider",
    "Prepayment remains unapplied after procedure completion",
    "Patient balance shows credit with unassigned provider"
]
for item in indicators:
    document.add_paragraph(item, style='List Bullet')

# --------------------------------------------------
# Section 2: Required Information Before Transfer
# --------------------------------------------------
document.add_heading('2. Required Information Before Transfer', level=2)
required_info = [
    "Patient name and ID",
    "Original payment date and amount",
    "Current provider assignment (source)",
    "Correct provider assignment (destination)",
    "Relevant procedure codes",
    "Reason for transfer",
    "Authorization (if required by amount threshold)"
]
for info in required_info:
    document.add_paragraph(info, style='List Bullet')

# --------------------------------------------------
# Section 3: Income Transfer Process Steps
# --------------------------------------------------
document.add_heading('3. Income Transfer Process Steps', level=2)
document.add_paragraph(
    "Note: This section will be enhanced with specific OpenDental UI screenshots and detailed "
    "step-by-step instructions once provided.", style='Intense Quote'
)

# A. Accessing the Income Transfer Function
document.add_heading('A. Accessing the Income Transfer Function', level=3)
access_steps = [
    "Log into OpenDental with appropriate user credentials",
    "Navigate to [Specific Path in OpenDental]",
    "Select the patient account requiring income transfer",
    "Verify patient information and account status"
]
for step in access_steps:
    document.add_paragraph(step, style='List Number')

# B. Executing the Transfer
document.add_heading('B. Executing the Transfer', level=3)
execute_steps = [
    "Locate the specific payment requiring transfer",
    "Select the income transfer function",
    "Document the source provider (current assignment)",
    "Select the destination provider (correct assignment)",
    "Verify the transfer amount",
    "Enter a standardized note (see format below)",
    "Complete the transfer"
]
for step in execute_steps:
    document.add_paragraph(step, style='List Number')

# C. Standardized Note Format
document.add_heading('C. Standardized Note Format', level=3)
document.add_paragraph("All income transfer notes must follow this format:")
document.add_paragraph(
    "```\nIncome Transfer: FROM [source provider] TO [destination provider] - [reason code] - [your initials]\n```",
    style='Intense Quote'
)
document.add_paragraph("Example:")
document.add_paragraph(
    "```\nIncome Transfer: FROM Unassigned TO Dr. Smith - PREPAY ALLOCATION - SW\n```",
    style='Intense Quote'
)
document.add_paragraph("Reason Codes:")
reason_codes = [
    "INITIAL ERROR: Payment was initially assigned incorrectly",
    "PREPAY ALLOCATION: Allocating prepayment to treating provider",
    "PROV CHANGE: Provider changed during treatment",
    "TX PLAN MOD: Treatment plan was modified",
    "OTHER: Other reason (requires additional explanation)"
]
for code in reason_codes:
    document.add_paragraph(code, style='List Bullet')

# --------------------------------------------------
# Section 4: Verification Process
# --------------------------------------------------
document.add_heading('4. Verification Process', level=2)
document.add_heading('Immediate Verification', level=3)
immediate_verification = [
    "Refresh the patient account",
    "Verify the payment now shows the correct provider",
    "Confirm the transfer note appears correctly",
    "Check that patient balance remains unchanged"
]
for step in immediate_verification:
    document.add_paragraph(step, style='List Number')

document.add_heading('Post-Transfer Verification', level=3)
post_verification = [
    "Run the daily split allocation report",
    "Verify provider production reports reflect the transfer",
    "Check for any unintended consequences (e.g., other allocations affected)"
]
for step in post_verification:
    document.add_paragraph(step, style='List Number')

# --------------------------------------------------
# Section 5: Common Issues and Troubleshooting
# --------------------------------------------------
document.add_heading('5. Common Issues and Troubleshooting', level=2)
issues = [
    ["Issue", "Potential Cause", "Solution"],
    ["Transfer not showing in reports", "Report parameters incorrect", "Adjust date ranges and filters"],
    ["Provider still shows as unassigned", "Transfer not completed properly", "Repeat transfer process"],
    ["Multiple transfers created", "Duplicate process execution", "Contact system administrator"],
    ["Transfer amount incorrect", "Split payment not fully selected", "Cancel and restart with correct amount"],
    ["System error during transfer", "Software or connection issue", "Document error, contact IT support"],
]
issue_table = document.add_table(rows=1, cols=3)
hdr = issue_table.rows[0].cells
hdr[0].text = issues[0][0]
hdr[1].text = issues[0][1]
hdr[2].text = issues[0][2]
for row in issues[1:]:
    cells = issue_table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]

# --------------------------------------------------
# Section 6: Approval Requirements
# --------------------------------------------------
document.add_heading('6. Approval Requirements', level=2)
approval = [
    ["Transfer Amount", "Approval Required"],
    ["< $500", "Self-verification"],
    ["$500 - $1,000", "Team lead verification"],
    ["> $1,000", "Manager approval required"],
    ["> $5,000", "Director/owner approval required"],
]
approval_table = document.add_table(rows=1, cols=2)
hdr = approval_table.rows[0].cells
hdr[0].text = approval[0][0]
hdr[1].text = approval[0][1]
for row in approval[1:]:
    cells = approval_table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]

# --------------------------------------------------
# Section 7: Documentation and Record Keeping
# --------------------------------------------------
document.add_heading('7. Documentation and Record Keeping', level=2)
doc_elements = [
    "Date and time of transfer",
    "Staff member performing transfer",
    "Source and destination providers",
    "Amount transferred",
    "Reason for transfer",
    "Approval (if required)",
    "Verification completion"
]
for element in doc_elements:
    document.add_paragraph(element, style='List Bullet')

# --------------------------------------------------
# Section 8: Monitoring and Analytics
# --------------------------------------------------
document.add_heading('8. Monitoring and Analytics', level=2)

# Daily Monitoring
document.add_heading('Daily Monitoring', level=3)
daily_points = [
    "Review all new income transfers",
    "Verify proper documentation",
    "Check for patterns requiring system or process changes"
]
for point in daily_points:
    document.add_paragraph(point, style='List Bullet')

# Weekly Reporting
document.add_heading('Weekly Reporting', level=3)
weekly_points = [
    "Total number and value of income transfers",
    "Transfers by reason code",
    "Transfers by staff member",
    "Provider impact analysis"
]
for point in weekly_points:
    document.add_paragraph(point, style='List Bullet')

# Monthly Analysis
document.add_heading('Monthly Analysis', level=3)
monthly_points = [
    "Identify root causes of transfers",
    "Assess staff training needs",
    "Evaluate process effectiveness",
    "Recommend system or procedure changes"
]
for point in monthly_points:
    document.add_paragraph(point, style='List Bullet')

# --------------------------------------------------
# Section 9: Training Requirements
# --------------------------------------------------
document.add_heading('9. Training Requirements', level=2)
training_steps = [
    "Complete initial income transfer training",
    "Demonstrate proficiency in the process",
    "Complete refresher training annually",
    "Review updated procedures as released"
]
for step in training_steps:
    document.add_paragraph(step, style='List Number')

# --------------------------------------------------
# Section 10: Audit Procedures
# --------------------------------------------------
document.add_heading('10. Audit Procedures', level=2)
audit_points = [
    "Randomly (10% of all transfers)",
    "For all transfers over $1,000",
    "For any provider with transfers exceeding 5% of monthly production",
    "As part of quarterly financial reviews"
]
for point in audit_points:
    document.add_paragraph(point, style='List Bullet')

# --------------------------------------------------
# Implementation Timeline
# --------------------------------------------------
document.add_heading('Implementation Timeline', level=2)
timeline = [
    ["Phase", "Duration", "Activities"],
    ["1: Training", "1 week", "Staff training on new procedures"],
    ["2: Supervised Operation", "2 weeks", "Transfers performed under supervision"],
    ["3: Audit Period", "1 month", "100% audit of all transfers"],
    ["4: Normal Operation", "Ongoing", "Regular process with standard audit rates"],
]
timeline_table = document.add_table(rows=1, cols=3)
hdr = timeline_table.rows[0].cells
hdr[0].text = timeline[0][0]
hdr[1].text = timeline[0][1]
hdr[2].text = timeline[0][2]
for row in timeline[1:]:
    cells = timeline_table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]

# --------------------------------------------------
# Attachments and Resources
# --------------------------------------------------
document.add_heading('Attachments and Resources', level=2)
document.add_paragraph("Note: This section will include:", style='List Bullet')
attachments = [
    "Screenshots of OpenDental UI for the transfer process",
    "Sample reports for verification",
    "Quick reference guide for common scenarios"
]
for attach in attachments:
    document.add_paragraph(attach, style='List Bullet')

# --------------------------------------------------
# Revision History
# --------------------------------------------------
document.add_heading('Revision History', level=2)
revision = [
    ["Version", "Date", "Changes", "Author"],
    ["1.0", "[Current Date]", "Initial document creation", "[Your Organization]"],
    ["1.1", "[Current Date]", "Added Unassigned Payments Analysis", "[Your Organization]"],
]
revision_table = document.add_table(rows=1, cols=4)
hdr = revision_table.rows[0].cells
hdr[0].text = revision[0][0]
hdr[1].text = revision[0][1]
hdr[2].text = revision[0][2]
hdr[3].text = revision[0][3]
for row in revision[1:]:
    cells = revision_table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]
    cells[3].text = row[3]

# --------------------------------------------------
# Approval Section
# --------------------------------------------------
document.add_paragraph("\nApproved by: ________________________")
document.add_paragraph("Date: _______________________________")

# --------------------------------------------------
# Appendix A: Unassigned Payments Analysis
# --------------------------------------------------
document.add_heading('Appendix A: Unassigned Payments Analysis', level=2)
document.add_heading('Current Issues Analysis', level=3)
document.add_paragraph(
    "Based on a comprehensive analysis of unassigned provider transactions, the following patterns have been identified:"
)

# Key Contributors
document.add_heading('Key Contributors', level=3)
document.add_paragraph("1. User Patterns:", style='List Number')
user_patterns = [
    "Sophie (16 transactions) is the primary contributor to unassigned payments",
    "Chelsea, Emily, and Melanie each contributed approximately 11 transactions",
    "The Admin Group department accounts for the highest number of unassigned transactions (32)"
]
for pattern in user_patterns:
    document.add_paragraph(pattern, style='List Bullet')

# Payment Types
document.add_heading('Payment Types', level=3)
payment_types = [
    "Credit Card transactions represent the majority of unassigned payments",
    "Patient Refund transactions are the second most common type",
    "Check and Care Credit transactions also appear in unassigned payments"
]
for pt in payment_types:
    document.add_paragraph(pt, style='List Bullet')

# Time Patterns
document.add_heading('Time Patterns', level=3)
time_patterns = [
    "February 15th shows a significant spike in unassigned transactions",
    "Tuesdays have the highest frequency of unassigned transaction entries",
    "Transaction frequency shows an irregular pattern suggesting batch processing"
]
for tp in time_patterns:
    document.add_paragraph(tp, style='List Bullet')

# Provider Associations
document.add_heading('Provider Associations', level=3)
provider_assoc = [
    "Dr. Timothy Kamp and Dr. Beau Schneider have the most patients with unassigned payments",
    "This suggests their patients' payments are being processed without proper provider assignment"
]
for pa in provider_assoc:
    document.add_paragraph(pa, style='List Bullet')

# Root Causes
document.add_heading('Root Causes', level=3)
root_causes = [
    "Training Gaps: Users, particularly Sophie, may not understand the importance of assigning providers to payments or the financial impact of unassigned payments.",
    "System Design Issues: The payment entry screen may not make the provider field mandatory or prominent enough during the payment entry process.",
    "Workflow Disconnects: Payments are likely being processed separately from treatment documentation, causing disconnection between provider information and the payment process.",
    "Time Pressure: The Tuesday pattern and transaction spikes suggest higher workloads may be leading to shortcuts in data entry.",
    "Credit Card Processing: The predominance of credit card transactions suggests a specific issue in how these payments are being processed in the system."
]
for cause in root_causes:
    document.add_paragraph(cause, style='List Number')

# Targeted Intervention Plan
document.add_heading('Targeted Intervention Plan', level=3)
document.add_heading('High Priority Actions', level=4)
high_priority = [
    "Individual Training: Provide targeted one-on-one training for Sophie focusing on proper provider assignment; observe Sophie's workflow; follow up with Chelsea and Emily.",
    "System Enhancement: Implement provider field validation; add warning messages for unassigned provider selection.",
    "Process Improvement: Create a daily report of unassigned provider transactions; implement a quick review process at the end of each day, especially for Tuesdays."
]
for action in high_priority:
    document.add_paragraph(action, style='List Bullet')

document.add_heading('Medium Priority Actions', level=4)
medium_priority = [
    "Department Training: Schedule department-wide training for the Admin Group on proper payment processing; develop visual step-by-step guides emphasizing provider assignment.",
    "Payment Type Controls: Implement specific validation controls for Credit Card payment processing; review the Patient Refund workflow.",
    "Workload Management: Review staffing and workflow on Tuesdays; consider redistributing tasks to reduce data entry burdens during peak times."
]
for action in medium_priority:
    document.add_paragraph(action, style='List Bullet')

document.add_heading('Monitoring Plan', level=3)
monitoring_plan = [
    "Daily Monitoring: Track Sophie's transactions daily for the first two weeks following training; review all unassigned provider transactions each morning.",
    "Weekly Analysis: Analyze unassigned payments by user, payment type, and day of week; provide feedback to staff on improvement trends.",
    "Monthly Review: Evaluate the effectiveness of interventions; adjust training and processes based on ongoing data."
]
for plan in monitoring_plan:
    document.add_paragraph(plan, style='List Number')

# --------------------------------------------------
# Save the Document
# --------------------------------------------------
# Create the docs directory if it doesn't exist
docs_dir = "scripts/validation/payment_split/docs"
os.makedirs(docs_dir, exist_ok=True)

# Save the document to the docs directory
output_path = os.path.join(docs_dir, "Income_Transfer_Workflow.docx")
document.save(output_path)
print(f"Document created successfully as '{output_path}'.")
